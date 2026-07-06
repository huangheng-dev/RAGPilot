import io

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook

from ragpilot_worker.domain import chunking
from ragpilot_worker.domain.chunking import normalize_text


def test_normalize_text_supports_plain_text() -> None:
    parsed_document = normalize_text(
        b"Hello\r\nWorld\r\n",
        content_type="text/plain",
        file_name="notes.txt",
    )

    assert parsed_document.parser_name == "plain_text_parser"
    assert parsed_document.text == "Hello\nWorld"


def test_normalize_text_supports_html() -> None:
    parsed_document = normalize_text(
        b"<html><body><h1>Title</h1><p>Hello <strong>world</strong>.</p></body></html>",
        content_type="text/html; charset=utf-8",
        file_name="overview.html",
    )

    assert parsed_document.parser_name == "html_parser"
    assert parsed_document.text == "Title\nHello world."


def test_normalize_text_supports_csv() -> None:
    parsed_document = normalize_text(
        "name,role\nAlice,Operator\nBob,Reviewer\n".encode("utf-8"),
        content_type="text/csv",
        file_name="members.csv",
    )

    assert parsed_document.parser_name == "csv_parser"
    assert parsed_document.text == "name | role\nAlice | Operator\nBob | Reviewer"


def test_normalize_text_supports_json() -> None:
    parsed_document = normalize_text(
        b'{"tenant":"ragpilot-demo","active":true}',
        content_type="application/json",
        file_name="scope.json",
    )

    assert parsed_document.parser_name == "json_parser"
    assert '"active": true' in parsed_document.text
    assert '"tenant": "ragpilot-demo"' in parsed_document.text


def test_normalize_text_supports_pdf(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePdfPage:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    class FakePdfReader:
        def __init__(self, stream: io.BytesIO) -> None:
            assert stream.read(5) == b"%PDF-"
            self.pages = [
                FakePdfPage("RAGPilot PDF handbook"),
                FakePdfPage("Durable ingestion workflow"),
            ]

    monkeypatch.setattr(chunking, "PdfReader", FakePdfReader)

    parsed_document = normalize_text(
        b"%PDF-1.7",
        content_type="application/pdf",
        file_name="handbook.pdf",
    )

    assert parsed_document.parser_name == "pdf_parser"
    assert parsed_document.text == "Page 1\nRAGPilot PDF handbook\n\nPage 2\nDurable ingestion workflow"


def test_normalize_text_supports_docx() -> None:
    document = DocxDocument()
    document.add_heading("RAGPilot Handbook", level=1)
    document.add_paragraph("Durable ingestion stays observable.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Owner"
    table.rows[0].cells[1].text = "Team"
    table.rows[1].cells[0].text = "Platform"
    table.rows[1].cells[1].text = "Operations"

    buffer = io.BytesIO()
    document.save(buffer)

    parsed_document = normalize_text(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_name="handbook.docx",
    )

    assert parsed_document.parser_name == "docx_parser"
    assert "RAGPilot Handbook" in parsed_document.text
    assert "Durable ingestion stays observable." in parsed_document.text
    assert "Table 1" in parsed_document.text
    assert "Platform | Operations" in parsed_document.text


def test_normalize_text_supports_xlsx() -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Summary"
    worksheet.append(["Metric", "Value"])
    worksheet.append(["Indexed documents", 12])
    worksheet.append(["Failed workflows", 2])

    buffer = io.BytesIO()
    workbook.save(buffer)
    workbook.close()

    parsed_document = normalize_text(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        file_name="summary.xlsx",
    )

    assert parsed_document.parser_name == "xlsx_parser"
    assert parsed_document.text == "Sheet: Summary\nMetric | Value\nIndexed documents | 12\nFailed workflows | 2"


def test_normalize_text_rejects_unsupported_types() -> None:
    with pytest.raises(ValueError, match="Unsupported document type for initial ingestion"):
        normalize_text(
            b"\x89PNG\r\n",
            content_type="image/png",
            file_name="diagram.png",
        )
