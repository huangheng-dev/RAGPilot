from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass
from datetime import datetime
from html import unescape

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


@dataclass(frozen=True)
class TextChunk:
    chunk_index: int
    content: str
    token_count: int
    metadata_json: dict[str, int | str | bool]


@dataclass(frozen=True)
class ParsedDocument:
    parser_name: str
    text: str


TEXT_CONTENT_TYPES = {"text/plain", "text/markdown"}
HTML_CONTENT_TYPES = {"text/html", "application/xhtml+xml"}
CSV_CONTENT_TYPES = {"text/csv", "application/csv"}
JSON_CONTENT_TYPES = {"application/json", "text/json"}
PDF_CONTENT_TYPES = {"application/pdf"}
DOCX_CONTENT_TYPES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
XLSX_CONTENT_TYPES = {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}
IMAGE_CONTENT_TYPES = {"image/png", "image/jpeg", "image/webp", "image/tiff", "image/bmp"}
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp")
MAX_OCR_IMAGE_PIXELS = 40_000_000


def normalize_text(content: bytes, *, content_type: str | None, file_name: str) -> ParsedDocument:
    lower_name = file_name.lower()
    normalized_content_type = _normalize_content_type(content_type)

    if normalized_content_type in TEXT_CONTENT_TYPES or lower_name.endswith((".txt", ".md", ".markdown")):
        return ParsedDocument(parser_name="plain_text_parser", text=_normalize_plain_text(content))
    if normalized_content_type in HTML_CONTENT_TYPES or lower_name.endswith((".html", ".htm")):
        return ParsedDocument(parser_name="html_parser", text=_normalize_html_text(content))
    if normalized_content_type in CSV_CONTENT_TYPES or lower_name.endswith(".csv"):
        return ParsedDocument(parser_name="csv_parser", text=_normalize_csv_text(content))
    if normalized_content_type in JSON_CONTENT_TYPES or lower_name.endswith(".json"):
        return ParsedDocument(parser_name="json_parser", text=_normalize_json_text(content))
    if normalized_content_type in PDF_CONTENT_TYPES or lower_name.endswith(".pdf"):
        extracted = _normalize_pdf_text(content)
        if extracted.strip():
            return ParsedDocument(parser_name="pdf_parser", text=extracted)
        return ParsedDocument(parser_name="pdf_ocr_parser", text=_normalize_pdf_ocr_text(content))
    if normalized_content_type in DOCX_CONTENT_TYPES or lower_name.endswith(".docx"):
        return ParsedDocument(parser_name="docx_parser", text=_normalize_docx_text(content))
    if normalized_content_type in XLSX_CONTENT_TYPES or lower_name.endswith(".xlsx"):
        return ParsedDocument(parser_name="xlsx_parser", text=_normalize_xlsx_text(content))
    if normalized_content_type in IMAGE_CONTENT_TYPES or lower_name.endswith(IMAGE_EXTENSIONS):
        return ParsedDocument(parser_name="image_ocr_parser", text=_normalize_image_ocr_text(content))

    raise ValueError(f"Unsupported document type for initial ingestion: {content_type or file_name}")


def _normalize_content_type(content_type: str | None) -> str | None:
    if not content_type:
        return None
    return content_type.split(";", 1)[0].strip().lower() or None


def _decode_text(content: bytes) -> str:
    return content.decode("utf-8-sig", errors="replace").replace("\r\n", "\n").replace("\r", "\n")


def _normalize_plain_text(content: bytes) -> str:
    return _decode_text(content).strip()


def _normalize_html_text(content: bytes) -> str:
    raw_html = _decode_text(content)
    without_scripts = re.sub(r"(?is)<(script|style)\b[^>]*>.*?</\1>", " ", raw_html)
    block_spaced = re.sub(r"(?i)</?(article|aside|blockquote|body|br|div|footer|h[1-6]|header|hr|li|main|ol|p|section|table|td|th|tr|ul)\b[^>]*>", "\n", without_scripts)
    without_tags = re.sub(r"(?is)<[^>]+>", " ", block_spaced)
    unescaped = unescape(without_tags)
    lines = [re.sub(r"\s+", " ", line).strip() for line in unescaped.splitlines()]
    normalized_text = "\n".join(line for line in lines if line).strip()
    return re.sub(r"\s+([,.:;!?])", r"\1", normalized_text)


def _normalize_csv_text(content: bytes) -> str:
    reader = csv.reader(io.StringIO(_decode_text(content)))
    rows: list[str] = []

    for row in reader:
        cleaned_cells = [re.sub(r"\s+", " ", cell).strip() for cell in row]
        if any(cleaned_cells):
            rows.append(" | ".join(cleaned_cells))

    return "\n".join(rows).strip()


def _normalize_json_text(content: bytes) -> str:
    payload = json.loads(_decode_text(content))
    return json.dumps(payload, indent=2, ensure_ascii=False, sort_keys=True).strip()


def _normalize_pdf_text(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    sections: list[str] = []

    for page_index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        normalized_page_text = _normalize_multiline_text(page_text)
        if normalized_page_text:
            sections.append(f"Page {page_index}\n{normalized_page_text}")

    return "\n\n".join(sections).strip()


def _normalize_pdf_ocr_text(content: bytes) -> str:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as error:
        raise ValueError("Scanned PDF requires the governed OCR runtime dependencies.") from error

    sections: list[str] = []
    document = fitz.open(stream=content, filetype="pdf")
    try:
        for page_index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
            page_text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            normalized = _normalize_multiline_text(page_text)
            if normalized:
                sections.append(f"Page {page_index} [OCR]\n{normalized}")
    finally:
        document.close()
    return "\n\n".join(sections).strip()


def _normalize_image_ocr_text(content: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image, UnidentifiedImageError
    except ImportError as error:
        raise ValueError("Image OCR requires the governed OCR runtime dependencies.") from error

    try:
        with Image.open(io.BytesIO(content)) as source_image:
            width, height = source_image.size
            if width <= 0 or height <= 0 or width * height > MAX_OCR_IMAGE_PIXELS:
                raise ValueError("Image dimensions exceed the governed OCR safety limit.")
            image = source_image.convert("RGB")
            image.load()
    except (UnidentifiedImageError, OSError) as error:
        raise ValueError("Image content is invalid or unsupported by the OCR runtime.") from error

    text = pytesseract.image_to_string(image, lang="chi_sim+eng")
    normalized = _normalize_multiline_text(text)
    return f"Image [OCR]\n{normalized}" if normalized else ""


def _normalize_docx_text(content: bytes) -> str:
    document = DocxDocument(io.BytesIO(content))
    sections: list[str] = []

    for paragraph in document.paragraphs:
        normalized_paragraph = _normalize_single_line_text(paragraph.text)
        if normalized_paragraph:
            sections.append(normalized_paragraph)

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            normalized_cells = [_normalize_single_line_text(cell.text) for cell in row.cells]
            if any(normalized_cells):
                rows.append(" | ".join(normalized_cells))

        if rows:
            sections.append(f"Table {table_index}\n" + "\n".join(rows))

    return "\n\n".join(sections).strip()


def _normalize_xlsx_text(content: bytes) -> str:
    workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    sections: list[str] = []

    try:
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                normalized_cells = [_normalize_spreadsheet_value(cell_value) for cell_value in row]
                if any(normalized_cells):
                    rows.append(" | ".join(normalized_cells))

            if rows:
                sections.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
    finally:
        workbook.close()

    return "\n\n".join(sections).strip()


def _normalize_multiline_text(value: str) -> str:
    lines = [_normalize_single_line_text(line) for line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _normalize_single_line_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value)).strip()


def _normalize_spreadsheet_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat()
    return _normalize_single_line_text(value)


def build_text_chunks(text: str, *, chunk_size: int, chunk_overlap: int) -> list[TextChunk]:
    if not text:
        return []

    collapsed_text = re.sub(r"\n{3,}", "\n\n", text)
    paragraphs = [paragraph.strip() for paragraph in collapsed_text.split("\n\n") if paragraph.strip()]
    chunks: list[TextChunk] = []
    current_content = ""
    current_start = 0
    current_locator: dict[str, int | str | bool] = {}
    cursor = 0

    for paragraph in paragraphs:
        paragraph_locator = _extract_source_locator(paragraph)
        if paragraph_locator and current_content:
            chunks.append(_build_chunk(len(chunks), current_content, current_start, current_locator))
            current_content = ""
            current_locator = {}
        if not current_content and paragraph_locator:
            current_locator = paragraph_locator
        paragraph_text = paragraph if not current_content else f"{current_content}\n\n{paragraph}"
        if len(paragraph_text) <= chunk_size:
            if not current_content:
                current_start = cursor
            current_content = paragraph_text
            cursor += len(paragraph) + 2
            continue

        if current_content:
            chunks.append(_build_chunk(len(chunks), current_content, current_start, current_locator))
            overlap_text = current_content[-chunk_overlap:] if chunk_overlap > 0 else ""
            current_content = overlap_text.strip()
            current_start = max(current_start, cursor - len(current_content))

        remaining = paragraph
        while len(remaining) > chunk_size:
            piece = remaining[:chunk_size]
            chunks.append(_build_chunk(len(chunks), piece, cursor, paragraph_locator or current_locator))
            remaining = remaining[max(chunk_size - chunk_overlap, 1):]
            cursor += max(chunk_size - chunk_overlap, 1)

        current_content = remaining
        current_locator = paragraph_locator or current_locator
        current_start = cursor
        cursor += len(paragraph) + 2

    if current_content:
        chunks.append(_build_chunk(len(chunks), current_content, current_start, current_locator))

    return chunks


def _build_chunk(
    chunk_index: int, content: str, start_char: int, locator: dict[str, int | str | bool] | None = None,
) -> TextChunk:
    normalized_content = content.strip()
    token_count = len(normalized_content.split())
    return TextChunk(
        chunk_index=chunk_index,
        content=normalized_content,
        token_count=token_count,
        metadata_json={
            "start_char": start_char,
            "end_char": start_char + len(normalized_content),
            **(locator or {}),
        },
    )


def _extract_source_locator(paragraph: str) -> dict[str, int | str | bool]:
    first_line = paragraph.splitlines()[0].strip() if paragraph else ""
    page_match = re.fullmatch(r"Page (\d+)( \[OCR\])?", first_line, flags=re.IGNORECASE)
    if page_match:
        page_number = int(page_match.group(1))
        is_ocr = bool(page_match.group(2))
        return {
            "source_location_type": "page",
            "source_location_label": f"Page {page_number}" + (" (OCR)" if is_ocr else ""),
            "source_page_number": page_number,
            "source_is_ocr": is_ocr,
        }
    sheet_match = re.fullmatch(r"Sheet: (.+)", first_line, flags=re.IGNORECASE)
    if sheet_match:
        sheet_name = sheet_match.group(1).strip()
        return {
            "source_location_type": "sheet",
            "source_location_label": f"Sheet: {sheet_name}",
            "source_sheet_name": sheet_name,
        }
    table_match = re.fullmatch(r"Table (\d+)", first_line, flags=re.IGNORECASE)
    if table_match:
        table_number = int(table_match.group(1))
        return {
            "source_location_type": "table",
            "source_location_label": f"Table {table_number}",
            "source_table_number": table_number,
        }
    if first_line.lower() == "image [ocr]":
        return {
            "source_location_type": "image",
            "source_location_label": "Image (OCR)",
            "source_is_ocr": True,
        }
    return {}
